import os
from datetime import datetime
import time
import pandas as pd
from io import BytesIO
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
import warnings
from sqlalchemy import create_engine
from sqlalchemy.sql import functions 
from sqlalchemy.orm import sessionmaker
from Endzone_Database.classes import Game, Formation
import logging
from Endzone_Reports.graph_utils import Graph_Utils

class PregameReport:
    def __init__(self, input: list, team_of_interest: str, team_code: str, type: str):
        self.jobId = str(int(time.time()))
        self.input = input
        self.team_code = team_code
        self.team_of_interest = team_of_interest
        self.type = type
        self.report = Document()

        # Logging
        logging.basicConfig(filename = os.path.join(os.getcwd(), "Endzone_Reports/Logs/PostGame_Log.txt"), level=logging.INFO)
        logging.info("\t Job " + self.jobId +  " Start | \t" + str(self.input) + "\t" + str(datetime.fromtimestamp(time.time()).isoformat()))

        self.data = self.GetData()
        self.data = self.PrepData()
        self.ReportManager()
    
    def GetData(self):
        db_engine = create_engine("postgresql://endzone:Hercules22!@172.104.25.168/Endzone")
        Session = sessionmaker(db_engine)
        session = Session()
        data = pd.read_sql(session.query(Game).filter(Game.Owner_Team_Code == self.team_code).filter(functions.concat(Game.Team_Name, " VS ", Game.Opponent_Name, " | ", Game.Year).in_(self.input)).statement, db_engine)
        df_forms = pd.read_sql(session.query(Formation).filter(Formation.Team_Code == self.team_code).statement, db_engine)
        data = pd.merge(data, df_forms, on='Formation', how='left')
        if len(data) > 0:
            return data
        else:
            logging.info("\t Job " + self.jobId +  " Error | \t" + str(self.input) + "\t" + str(datetime.fromtimestamp(time.time()).isoformat()))
            logging.error("No Data found")
            return

    def PrepData(self):
        self.data["DownGroup"] = ""
        self.data["FieldGroup"] = ""
        self.data["Pressure_Stunt"] = ""
        self.data["Pressure_Center"] = ""
        self.data["Pressure_Edge"] = ""
        self.data["Pressure_Existence"] = ""
        self.data["Number_Rushers"] = 0
        self.data["Coverage_Type"] = ""
        self.data["Run_Type"] = ""
        self.data["Personal"] = ""
        for index, row in self.data.iterrows():
            try:
                self.data.at[index, "Number_Rushers"] = str(self.data.at[index, "D_Formation"])[0:1]
                if self.data.at[index, "Distance"] <= 3:
                    self.data.at[index, "DownGroup"] = "Short"
                elif self.data.at[index, "Distance"] > 3 and self.data.at[index, "Distance"] <= 6:
                    self.data.at[index, "DownGroup"] = "Medium"
                elif self.data.at[index, "Distance"] > 6:
                    self.data.at[index, "DownGroup"] = "Long"
                if self.data.at[index, "Yard"] <= 33:
                    self.data.at[index, "FieldGroup"] = "Backed Up"
                elif self.data.at[index, "Yard"] > 33 and self.data.at[index, "Yard"] <= 66:
                    self.data.at[index, "FieldGroup"] = "Midfield"
                elif self.data.at[index, "Yard"] > 66:
                    self.data.at[index, "FieldGroup"] = "Scoring Position"
                if int(self.data.at[index, "Pressure_Middle"]) > 0:
                    self.data.at[index, "Pressure_Center"] = "Yes"
                    self.data.at[index, "Number_Rushers"] = int(self.data.at[index, "Number_Rushers"]) + int(self.data.at[index, "Pressure_Middle"])
                else:
                    self.data.at[index, "Pressure_Center"] = "No"
                if int(self.data.at[index, "Pressure_Right"]) > 0:
                    self.data.at[index, "Number_Rushers"] = int(self.data.at[index, "Number_Rushers"]) + int(self.data.at[index, "Pressure_Right"])
                if int(self.data.at[index, "Pressure_Left"]) > 0:
                    self.data.at[index, "Number_Rushers"] = int(self.data.at[index, "Number_Rushers"]) + int(self.data.at[index, "Pressure_Left"])
                if int(self.data.at[index, "Pressure_Left"]) > 0 or int(self.data.at[index, "Pressure_Right"]) > 0:
                    self.data.at[index, "Pressure_Edge"] = "Yes"
                else:
                    self.data.at[index, "Pressure_Edge"] = "No"
                if self.data.at[index, "Number_Rushers"] != str(self.data.at[index, "D_Formation"])[0:1]:
                    self.data.at[index, "Pressure_Existence"] = "Yes"
                else:
                    self.data.at[index, "Pressure_Existence"] = "No"       
                if "Zone" in str(self.data.at[index, "Coverage"]):
                    self.data.at[index, "Coverage_Type"] = "Zone"
                elif "Man" in str(self.data.at[index, "Coverage"]):
                    self.data.at[index, "Coverage_Type"] = "Man"
                else:
                    self.data.at[index, "Coverage_Type"] = "Unknown" 
                if "Outside" in self.data.at[index, "Play_Type"]:
                    self.data.at[index, "Run_Type"] = "Outside"
                elif "Inside" in self.data.at[index, "Play_Type"]:
                    self.data.at[index, "Run_Type"] = "Inside"
                else:
                    self.data.at[index, "Run_Type"] = "Null" 
            except Exception as e:
                logging.error(e)
                pass
        self.data["FieldSpot"] = self.data["Hash"] + "-" + self.data["FieldGroup"]
        self.data["Personnel"] = self.data["Wide_Recievers"].astype(str) + "WR | " + self.data["Tight_Ends"].astype(str) + "TE | " + self.data["Running_Backs"].astype(str) + "RB"
        return self.data
    
    def ReportManager(self):
        try:
            ### Title Page
            title = self.report.add_heading("EVFB Analytics Pregame Analysis")
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph = self.report.add_paragraph()
            paragraph.add_run("Opponents in Data: ")
            for team in self.data.Possession.unique():
                if team not in self.team_of_interest:
                    paragraph.add_run(team + ", ")
            for team in self.data.Possession.unique():
                    if team in self.team_of_interest:
                        if self.type == "Defense" or self.type == "All":
                            ### Title Page - Table of Contents
                            self.report.add_heading('%s - Offense' %team, 2)
                            self.report.add_paragraph('Efficiencies', style='List Number 2')
                            self.report.add_paragraph('Yardage', style='List Number 2')
                            self.report.add_paragraph('Play Type', style='List Number 2')
                            self.report.add_paragraph('Ball Carrier', style='List Number 2')
                            self.report.add_paragraph('1st Down', style='List Number 2')
                            self.report.add_paragraph('2nd Down', style='List Number 2')
                            self.report.add_paragraph('3rd Down', style='List Number 2')
                            self.report.add_paragraph('Pass Zones', style='List Number 2')
                            self.report.add_paragraph('Ideal Passing Coverage', style='List Number 2')
                            self.report.add_paragraph('Targets', style='List Number 2')
                            self.report.add_paragraph('Redzone', style='List Number 2')
                            self.report.add_paragraph('Boundary', style='List Number 2')
                            self.report.add_page_break()
                            DefenseSection(self.data, self.report, self.team_of_interest)
                        if self.type == "Offense" or self.type == "All":
                            self.report.add_heading('%s - Defense' %team, 2)
                            self.report.add_heading('Early mockup of EVFB Analytics report. Will be rewritten in next week. Request features as you want', 4)
                            self.report.add_paragraph('Efficiencies -- IN TEST', style='List Number 2')
                            self.report.add_paragraph('Yardage -- IN TEST', style='List Number 2')
                            self.report.add_paragraph('Rushers -- IN TEST', style='List Number 2')
                            self.report.add_paragraph('Coverage -- IN TEST', style='List Number 2')
                            self.report.add_paragraph('Pressure -- IN TEST', style='List Number 2')
                            self.report.add_paragraph('Boundary -- IN TEST', style='List Number 2')
                            self.report.add_paragraph('1st Down -- IN DEV', style='List Number 2')
                            self.report.add_paragraph('2nd Down -- IN DEV', style='List Number 2')
                            self.report.add_paragraph('3rd Down -- IN DEV', style='List Number 2')
                            self.report.add_paragraph('Redzone -- IN DEV', style='List Number 2')
                            self.report.add_paragraph('Weakness Analysis -- IN DEV', style='List Number 2')
                            self.report.add_page_break()
                            OffenseSection(self.data, self.report, self.team_of_interest)
                            pass


            self.output_path = os.path.dirname(__file__) + "/Outputs/pre_game/EVFB_PreGameReport_%s.docx" %self.jobId
            self.report.save(self.output_path)
            logging.info("\t Job " + self.jobId +  " Complete | \t" + str(self.input) + "\t" + str(datetime.fromtimestamp(time.time()).isoformat()))
        except Exception as e:
             print(e)
             logging.error(e)
             return

class OffenseSection(Graph_Utils):
    def __init__(self, data, report, team_of_interest, **kwargs):
        self.data = data
        self.report = report
        self.team_of_interest = team_of_interest
        self.memory_space = BytesIO()
        self.ReportManager()


    def ReportManager(self):
        self.GetOData()
        self.Effenciies()
        self.Yards()
        self.Rushers()
        self.Coverage()
        self.Pressure()
        self.Boundary()
        return

    def GetOData(self):
        self.o_report = self.data[(self.data["Possession"] != self.team_of_interest)]
        self.o_report["Distance_Hash"] = self.o_report["Hash"] + " - " + self.o_report["DownGroup"]
        self.o_passing_zones = self.o_report[self.o_report["Pass_Zone"] != "Non Passing Play"]
        self.o_runs = self.o_report[(self.o_report["Play_Type"] == "Inside Run") | (self.o_report["Play_Type"] == "Outside Run")]
        self.o_pass = self.o_report[(self.o_report["Play_Type"] == "Boot Pass") | (self.o_report["Play_Type"] == "Pocket Pass")]
        self.o_1_down = self.o_report[self.o_report["Down"] == 1]
        self.o_2_down = self.o_report[self.o_report["Down"] == 2]
        self.o_3_down = self.o_report[self.o_report["Down"] == 3]
        self.o_redzone = self.o_report[self.o_report["Yard"] >= 80]
        self.o_boundary = self.o_report[(self.o_report["Hash"] == "Right") | (self.o_report["Hash"] == "Left")]

        ## Boundary Prep
        self.o_boundary["Boundary"] = " "
        self.o_boundary["Boundary_Pressure"] = " "
        self.o_boundary["Wide_Pressure"] = " "

        for index, row in self.o_boundary.iterrows():
            if self.o_boundary.at[index, "Hash"] == self.o_boundary.at[index, "Play_Type_Dir"]:
                self.o_boundary.at[index, "Boundary"] = "Into Slideline"
            else:
                self.o_boundary.at[index, "Boundary"] = "Away from Sideline"

            if self.o_boundary.at[index, "Hash"] == "Left" and self.o_boundary.at[index, "Pressure_Left"] > 0:
                self.o_boundary.at[index, "Boundary_Pressure"] = "From Boundary"
            elif self.o_boundary.at[index, "Hash"] == "Right" and self.o_boundary.at[index, "Pressure_Right"] > 0:
                self.o_boundary.at[index, "Boundary_Pressure"] = "From Boundary"
            else:
                self.o_boundary.at[index, "Boundary_Pressure"] = "No Boundary Pressure"

            if self.o_boundary.at[index, "Hash"] == "Left" and self.o_boundary.at[index, "Pressure_Right"] > 0:
                self.o_boundary.at[index, "Wide_Pressure"] = "From Wideside"
            elif self.o_boundary.at[index, "Hash"] == "Right" and self.o_boundary.at[index, "Pressure_Left"] > 0:
                self.o_boundary.at[index, "Wide_Pressure"] = "From Wideside"
            else:
                self.o_boundary.at[index, "Wide_Pressure"] = "No Wideside Pressure"

        ### Redzone Data Prep
        self.o_redzone["Redzone"] = " " 
        for index, row in self.o_redzone.iterrows():
            if self.o_redzone.at[index, "Yard"] >= 80 and self.o_redzone.at[index, "Yard"] <= 90:
                self.o_redzone.at[index, "Redzone"] = "Yard Marker 20 - 10"
            elif self.o_redzone.at[index, "Yard"] > 90 and self.o_redzone.at[index, "Yard"] <= 95:
                self.o_redzone.at[index, "Redzone"] = "Yard Marker 10 - 5"
            elif self.o_redzone.at[index, "Yard"] > 95:
                self.o_redzone.at[index, "Redzone"] = "Yard Marker 5 - Endzone"
        return self

    def Effenciies(self):
        ## Set Headers
        table = self.report.add_table(10, 2)
        table.style = 'Light Grid'
        headers = table.rows[0].cells
        headers[0].text = "%s Offensive Statisics" %self.team_of_interest
        headers[1].text = "Total Plays: " + str(len(self.o_report))
        ### Make Bold
        table.rows[0].cells[0].paragraphs[0].runs[0].font.bold = True
        table.rows[0].cells[1].paragraphs[0].runs[0].font.bold = True
        ## Cleanup Variable
        del headers
        cells = table.rows[1].cells
        cells[0].text = "Overall Offensive Efficiency (Gain less than 3)"
        try:    
            cells[1].text = str(round(len(self.o_report.query('Result <= 3')) / len(self.o_report) * 100, 2)) + "%" + " | Observations: " + str(len(self.o_report))
        except: 
            cells[1].text = "No Data"

        cells = table.rows[2].cells
        cells[0].text = "Overall Offensive Efficiency vs Inside Run (Gain less than 3)"
        try:
            cells[1].text = str(round(len(self.o_report.query('Result <= 3 & Play_Type == "Inside Run"' )) / len(self.o_report.query('Play_Type == "Inside Run"')) * 100, 2)) + "%" + " | Observations: " + str(len(self.o_report.query('Play_Type == "Inside Run"')))
        except: 
            cells[1].text = "No Data"

        cells = table.rows[3].cells
        cells[0].text = "Overall Offensive Efficiency vs Outside Run (Gain less than 3)"
        try:
            cells[1].text = str(round(len(self.o_report.query('Result <= 3 & Play_Type == "Outside Run"' )) / len(self.o_report.query('Play_Type == "Inside Run"')) * 100, 2)) + "%" + " | Observations: " + str(len(self.o_report.query('Play_Type == "Inside Run"')))
        except: 
            cells[1].text = "No Data"

        cells = table.rows[4].cells
        cells[0].text = "Overall Offensive Efficiency vs Pocket Pass (Gain less than 3)"
        try:
            cells[1].text = str(round(len(self.o_report.query('Result <= 3 & Play_Type == "Pocket Pass"' )) / len(self.o_report.query('Play_Type == "Pocket Pass"')) * 100, 2)) + "%" + " | Observations: " + str(len(self.o_report.query('Play_Type == "Pocket Pass"')))
        except:
            cells[1].text = "No Data"

        cells = table.rows[5].cells
        cells[0].text = "Overall Offensive Efficiency vs Boot Pass (Gain less than 3)"
        try:
            cells[1].text = str(round(len(self.o_report.query('Result <= 3 & Play_Type == "Boot Pass"' )) / len(self.o_report.query('Play_Type == "Boot Pass"')) * 100, 2)) + "%" + " | Observations: " + str(len(self.o_report.query('Play_Type == "Boot Pass"')))
        except:
            cells[1].text = "No Data"

        cells = table.rows[6].cells
        cells[0].text = "Passing"
        table.rows[6].cells[0].paragraphs[0].runs[0].font.bold = True

        cells = table.rows[7].cells
        cells[0].text = "Allowed Completion Percentage"
        try:
            cells[1].text = str(round(len(self.o_pass.query('Pass_Zone != "Not Selected" & Result > 0' )) / len(self.o_pass.query('Pass_Zone != "Not Selected"')) * 100, 2)) + "%" + " | Observations: " + str(len(self.o_report.query('Pass_Zone == "Not Selected"')))
        except:
            cells[1].text = "No Data"

        cells = table.rows[8].cells
        cells[0].text = "Sack Rate"
        try:
            cells[1].text = str(round(len(self.o_pass.query('Result < 0 & Pass_Zone == "Non Passing Play"')) / len(self.o_pass) * 100, 2)) + "%" + " | Observations: " + str(len(self.o_pass))
        except: 
            cells[1].text = "No Data"

        cells = table.rows[9].cells
        cells[0].text = "3rd Down Conversion Rate"
        try:
            cells[1].text = str(round(len(self.o_report.query('Down == 3 & Result >= Distance')) / len(self.o_report.query('Down == 3')) * 100, 2)) + "%" + " | Observations: " + str(len(self.o_report.query('Down == 3')))
        except: 
            cells[1].text = "No Data"

        self.report.add_page_break()
        pass
    
    def Yards(self):
        title = self.report.add_heading('%s Defense - Yardage Allowed Breakdown' %self.team_of_interest, 2)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        ## Set Headers
        table = self.report.add_table(8, 5)
        table.style = 'Light Grid'
        headers = table.rows[0].cells
        headers[0].text = "Type"
        headers[1].text = "Total Yards"
        headers[2].text = "Yards Per Game"
        headers[3].text = "Yards Per Drive"
        headers[4].text = "Yards Per Attempt"

        ### Make Bold
        table.rows[0].cells[0].paragraphs[0].runs[0].font.bold = True
        table.rows[0].cells[1].paragraphs[0].runs[0].font.bold = True
        table.rows[0].cells[2].paragraphs[0].runs[0].font.bold = True
        table.rows[0].cells[3].paragraphs[0].runs[0].font.bold = True
        table.rows[0].cells[4].paragraphs[0].runs[0].font.bold = True

        unique_games = len(pd.concat([self.o_report['Team_Name'],self.o_report['Opponent_Name'],self.o_report['Year']]).unique())
        unique_drives = len(self.o_report["Drive"].unique())
        unqiue_plays = len(self.o_report)

        cells = table.rows[1].cells
        cells[0].text = "Total Yards"
        try:
            cells[1].text = str(self.o_report["Result"].sum())
            cells[2].text = '{:0.2f}'.format(self.o_report["Result"].sum() / unique_games)
            cells[3].text = '{:0.2f}'.format(self.o_report["Result"].sum() / unique_drives)
            cells[4].text = '{:0.2f}'.format(self.o_report["Result"].sum() /unqiue_plays)
        except:
            cells[1].text = "No Data"
            cells[2].text = "No Data"
            cells[3].text = "No Data"
            cells[4].text = "No Data"
    
        cells = table.rows[2].cells
        cells[0].text = "Rushing Yards"
        try:   
            cells[1].text = str(self.o_runs["Result"].sum())
            cells[2].text = '{:0.2f}'.format(self.o_runs["Result"].sum() / unique_games)
            cells[3].text = '{:0.2f}'.format(self.o_runs["Result"].sum() / unique_drives)
            cells[4].text = '{:0.2f}'.format(self.o_runs["Result"].sum() /unqiue_plays) 
        except: 
            cells[1].text = "No Data"
            cells[2].text = "No Data"
            cells[3].text = "No Data"
            cells[4].text = "No Data"    

        cells = table.rows[3].cells
        cells[0].text = "Inside Run Yards"
        try: 
            cells[1].text = str(self.o_report.query("Play_Type == 'Inside Run'")["Result"].sum())
            cells[2].text = '{:0.2f}'.format(self.o_report.query("Play_Type == 'Inside Run'")["Result"].sum() / unique_games)
            cells[3].text = '{:0.2f}'.format(self.o_report.query("Play_Type == 'Inside Run'")["Result"].sum() / unique_drives)
            cells[4].text = '{:0.2f}'.format(self.o_report.query("Play_Type == 'Inside Run'")["Result"].sum() /unqiue_plays) 
        except: 
            cells[1].text = "No Data"
            cells[2].text = "No Data"
            cells[3].text = "No Data"
            cells[4].text = "No Data"

        cells = table.rows[4].cells
        cells[0].text = "Outside Run Yards"
        try:    
            cells[1].text = str(self.o_report.query("Play_Type == 'Outside Run'")["Result"].sum())
            cells[2].text = '{:0.2f}'.format(self.o_report.query("Play_Type == 'Outside Run'")["Result"].sum() / unique_games)
            cells[3].text = '{:0.2f}'.format(self.o_report.query("Play_Type == 'Outside Run'")["Result"].sum() / unique_drives)
            cells[4].text = '{:0.2f}'.format(self.o_report.query("Play_Type == 'Outside Run'")["Result"].sum() /unqiue_plays) 
        except:
            cells[1].text = "No Data"
            cells[2].text = "No Data"
            cells[3].text = "No Data"
            cells[4].text = "No Data"

        cells = table.rows[5].cells
        cells[0].text = "Passing Yards"
        try:    
            cells[1].text = str(self.o_pass["Result"].sum())
            cells[2].text = '{:0.2f}'.format(self.o_pass["Result"].sum() / unique_games)
            cells[3].text = '{:0.2f}'.format(self.o_pass["Result"].sum() / unique_drives)
            cells[4].text = '{:0.2f}'.format(self.o_pass["Result"].sum() /unqiue_plays)
        except: 
            cells[1].text = "No Data"
            cells[2].text = "No Data"
            cells[3].text = "No Data"
            cells[4].text = "No Data"  

        cells = table.rows[6].cells
        cells[0].text = "Pocket Pass Yards"
        try:    
            cells[1].text = str(self.o_pass.query("Play_Type == 'Pocket Pass'")["Result"].sum())
            cells[2].text = '{:0.2f}'.format(self.o_pass.query("Play_Type == 'Pocket Pass'")["Result"].sum() / unique_games)
            cells[3].text = '{:0.2f}'.format(self.o_pass.query("Play_Type == 'Pocket Pass'")["Result"].sum() / unique_drives)
            cells[4].text = '{:0.2f}'.format(self.o_pass.query("Play_Type == 'Pocket Pass'")["Result"].sum() /unqiue_plays) 
        except: 
            cells[1].text = "No Data"
            cells[2].text = "No Data"
            cells[3].text = "No Data"
            cells[4].text = "No Data" 

        cells = table.rows[7].cells
        cells[0].text = "Boot Pass Yards"
        try:    

            cells[1].text = str(self.o_pass.query("Play_Type == 'Boot Pass'")["Result"].sum())
            cells[2].text = '{:0.2f}'.format(self.o_pass.query("Play_Type == 'Boot Pass'")["Result"].sum() / unique_games)
            cells[3].text = '{:0.2f}'.format(self.o_pass.query("Play_Type == 'Boot Pass'")["Result"].sum() / unique_drives)
            cells[4].text = '{:0.2f}'.format(self.o_pass.query("Play_Type == 'Boot Pass'")["Result"].sum() /unqiue_plays) 
        except: 
            cells[1].text = "No Data"
            cells[2].text = "No Data"
            cells[3].text = "No Data"
            cells[4].text = "No Data" 
        self.report.add_paragraph()

        self.Table(self.o_report.groupby(["Formation"]).agg({"Result": ["count", "sum", "mean", "median", "max"]}), "Play Type Yardage Allowed")
        self.Table(self.o_report.groupby(["Personnel"]).agg({"Result": ["count", "sum", "mean", "median", "max"]}), "Play Type Yardage Allowed")
        self.report.add_page_break()

    def Rushers(self):
        title = self.report.add_heading('%s Defense - Rush Analysis' %self.team_of_interest, 2)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.o_report["Down"] = self.o_report["Down"].astype(str)
        self.o_report["Number_Rushers"] = self.o_report["Number_Rushers"].astype(str)
        # Number of Rushers by Personnel and Formation\
        self.CrosstabDisplay(pd.crosstab(self.o_report.Formation, self.o_report.Number_Rushers, margins=True, normalize="index"), "Rushers by Formation")
        self.CrosstabDisplay(pd.crosstab(self.o_report.Personnel, self.o_report.Number_Rushers, margins=True, normalize="index"), "Rushers by Personnel")


        self.BarGraph(self.o_report, "Down", "Number_Rushers", "Down", "Number of Rushers", 6, 3.5)
        self.SwarmPlot(self.o_report, "Distance", "Number_Rushers", "Distance", "Number of Rushers", 6, 3.5)

        self.CrosstabDisplay(pd.crosstab(self.o_report.FieldGroup, self.o_report.Number_Rushers, margins=True, normalize="index"), "Rushers by Field Position")

        self.report.add_page_break()
        
    def Coverage(self):
        title = self.report.add_heading('%s Defense - Coverage' %self.team_of_interest, 2)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        self.CrosstabDisplay(pd.crosstab(self.o_report.Personnel, self.o_report.Coverage, margins=True, normalize="index"), "Coverage by Personnel")
        self.CrosstabDisplay(pd.crosstab(self.o_report.Formation, self.o_report.Coverage, margins=True, normalize="index"), "Coverage by Formation")

        self.BarGraph(self.o_report, "Down", "Coverage", "Down", "Coverage", 6, 3.5)
        self.SwarmPlot(self.o_report, "Distance", "Coverage", "Distance", "Coverage", 6, 3.5)

        self.CrosstabDisplay(pd.crosstab(self.o_report.FieldGroup, self.o_report.Coverage, margins=True, normalize="index"), "Coverage by Field Position")
        self.report.add_page_break()
        pass
    
    def Pressure(self):
        title = self.report.add_heading('%s Defense - Pressure (Blitz) Analysis' %self.team_of_interest, 2)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        self.o_report["Down"] = self.o_report["Down"].astype(str)
        self.o_report["Pressure_Existence"] = self.o_report["Pressure_Existence"].astype(str)
        self.o_report["Pressure_Edge"] = self.o_report["Pressure_Edge"].astype(str)
        self.o_report["Pressure_Middle"] = self.o_report["Pressure_Middle"].astype(str)

        self.CrosstabDisplay(pd.crosstab(self.o_report.Personnel, self.o_report.Pressure_Existence, margins=True, normalize="index"), "Blitz Rate by Formation")
        self.CrosstabDisplay(pd.crosstab(self.o_report.Formation, self.o_report.Pressure_Existence, margins=True, normalize="index"), "Blitz Rate by Personnel")
        self.BarGraph(self.o_report, "Down", "Pressure_Existence", "Down", "Blitz Rate by Down", 6, 3.5)
        self.SwarmPlot(self.o_report, "Distance", "Pressure_Existence", "Distance", "Blitz Rate by Distance", 6, 3.5)
        self.report.add_page_break()

        # Edge Pressure
        title = self.report.add_heading('%s Defense - Pressure Edge (Blitz) Analysis' %self.team_of_interest, 2)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.CrosstabDisplay(pd.crosstab(self.o_report.Personnel, self.o_report.Pressure_Edge, margins=True, normalize="index"), "Edge Blitz Rate by Formation")
        self.CrosstabDisplay(pd.crosstab(self.o_report.Formation, self.o_report.Pressure_Edge, margins=True, normalize="index"), "Edge Blitz Rate by Personnel")
        self.BarGraph(self.o_report, "Down", "Pressure_Edge", "Down", "Edge Blitz Rate by Down", 6, 3.5)
        self.SwarmPlot(self.o_report, "Distance", "Pressure_Edge", "Distance", "Edge Blitz Rate by Distance", 6, 3.5)
        self.report.add_page_break()

        # Middle Pressure
        title = self.report.add_heading('%s Defense - Pressure Middle (Blitz) Analysis' %self.team_of_interest, 2)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.CrosstabDisplay(pd.crosstab(self.o_report.Personnel, self.o_report.Pressure_Center, margins=True, normalize="index"), "Middle Blitz Rate by Formation")
        self.CrosstabDisplay(pd.crosstab(self.o_report.Formation, self.o_report.Pressure_Center, margins=True, normalize="index"), "Middle Blitz Rate by Personnel")
        self.BarGraph(self.o_report, "Down", "Pressure_Center", "Down", "Middle Blitz Rate by Down", 6, 3.5)
        self.SwarmPlot(self.o_report, "Distance", "Pressure_Center", "Distance", "Middle Blitz Rate by Distance", 6, 3.5)
        self.report.add_page_break()
    
    def Boundary(self):
        title = self.report.add_heading('%s Defense - Boundary Pressure Analysis' %self.team_of_interest, 2)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        self.o_boundary["Down"] = self.o_boundary["Down"].astype(str)
        # Boundary Pressure
        self.CrosstabDisplay(pd.crosstab(self.o_boundary.Hash, self.o_boundary.Boundary_Pressure, margins=True, normalize="index"), "Boundary Blitz Rate by Hash")
        self.CrosstabDisplay(pd.crosstab(self.o_boundary.Personnel, self.o_boundary.Boundary_Pressure, margins=True, normalize="index"), "Boundary Blitz Rate by Personnel")
        self.CrosstabDisplay(pd.crosstab(self.o_boundary.Down, self.o_boundary.Boundary_Pressure, margins=True, normalize="index"), "Boundary Blitz Rate by Down")
        self.report.add_page_break()

        # Boundary Coverage
        title = self.report.add_heading('%s Defense - Boundary Coverage Analysis' %self.team_of_interest, 2)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.CrosstabDisplay(pd.crosstab(self.o_boundary.Hash, self.o_boundary.Coverage, margins=True, normalize="index"), "Boundary Coverage by Hash")
        self.CrosstabDisplay(pd.crosstab(self.o_boundary.Personnel, self.o_boundary.Coverage, margins=True, normalize="index"), "Boundary Coverage by Personnel")
        self.CrosstabDisplay(pd.crosstab(self.o_boundary.Down, self.o_boundary.Coverage, margins=True, normalize="index"), "Boundary Coverage by Down")
        self.report.add_page_break()

        # Boundary Yards Rate
        title = self.report.add_heading('%s Defense - Boundary Yardage Analysis' %self.team_of_interest, 2)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.Table(self.o_boundary.groupby(["Play_Type", "Boundary"]).agg({"Result": ["count", "sum", "mean", "median", "max"]}), "Play Type Yardage Allowed")
        self.Table(self.o_boundary.groupby(["Hash", "Boundary"]).agg({"Result": ["count", "sum", "mean", "median", "max"]}), "Play Type Yardage Allowed")

    def DownDetail():
        pass
    
    def Redzone():
        pass
    
    def Temporal():
        pass

class DefenseSection(Graph_Utils):
    def __init__(self, data, report, team_of_interest):
        self.data = data
        self.report = report
        self.team_of_interest = team_of_interest
        self.memory_space = BytesIO()
        self.ReportManager()

    def ReportManager(self):
        self.GetData()
        self.Effenciies()
        self.Yards()
        self.PlayType()
        self.BallCarrier()
        self.Down()
        self.PassZone()
        self.IdealCoverage()
        self.Targets()
        self.Redzone()
        self.Boundary()
        self.Temporal()
        return

    def GetData(self):
        self.d_report = self.data[self.data["Possession"] == self.team_of_interest]
        self.d_report["Distance_Hash"] = self.d_report["Hash"] + " - " + self.d_report["DownGroup"]
        self.d_passing_zones = self.d_report[self.d_report["Pass_Zone"] != "Non Passing Play"]
        self.d_runs = self.d_report[(self.d_report["Play_Type"] == "Inside Run") | (self.d_report["Play_Type"] == "Outside Run")]
        self.d_pass = self.d_report[(self.d_report["Play_Type"] == "Boot Pass") | (self.d_report["Play_Type"] == "Pocket Pass")]
        self.d_1_down = self.d_report[self.d_report["Down"] == 1]
        self.d_2_down = self.d_report[self.d_report["Down"] == 2]
        self.d_3_down = self.d_report[self.d_report["Down"] == 3]
        self.d_redzone = self.d_report[self.d_report["Yard"] >= 80]
        self.d_boundary = self.d_report[(self.d_report["Hash"] == "Right") | (self.d_report["Hash"] == "Left")]
        ## Boundary Prep
        self.d_boundary["Boundary"] = " "
        self.d_boundary["Boundary_Pressure"] = " "
        self.d_boundary["Wide_Pressure"] = " "

        for index, row in self.d_boundary.iterrows():
            if self.d_boundary.at[index, "Hash"] == self.d_boundary.at[index, "Play_Type_Dir"]:
                self.d_boundary.at[index, "Boundary"] = "Into Slideline"
            else:
                self.d_boundary.at[index, "Boundary"] = "Away from Sideline"



        ### Redzone Data Prep
        self.d_redzone["Redzone"] = " " 
        for index, row in self.d_redzone.iterrows():
            if self.d_redzone.at[index, "Yard"] >= 80 and self.d_redzone.at[index, "Yard"] <= 90:
                self.d_redzone.at[index, "Redzone"] = "Yard Marker 20 - 10"
            elif self.d_redzone.at[index, "Yard"] > 90 and self.d_redzone.at[index, "Yard"] <= 95:
                self.d_redzone.at[index, "Redzone"] = "Yard Marker 10 - 5"
            elif self.d_redzone.at[index, "Yard"] > 95:
                self.d_redzone.at[index, "Redzone"] = "Yard Marker 5 - Endzone"
        return self

    def Effenciies(self):
        ## Set Headers
        table = self.report.add_table(18, 2)
        table.style = 'Light Grid'
        headers = table.rows[0].cells
        headers[0].text = "%s Offensive Statisics" %self.team_of_interest
        headers[1].text = "Total Plays: " + str(len(self.d_report))
        ### Make Bold
        table.rows[0].cells[0].paragraphs[0].runs[0].font.bold = True
        table.rows[0].cells[1].paragraphs[0].runs[0].font.bold = True
        ## Cleanup Variable
        del headers
        cells = table.rows[1].cells
        cells[0].text = "Overall Offensive Efficiency (Gain greater than 3)"
        try:    
            cells[1].text = str(round(len(self.d_report.query('Result >= 3')) / len(self.d_report) * 100, 2)) + "%" + " | Observations: " + str(len(self.d_report))
        except: 
            cells[1].text = "No Data"

        cells = table.rows[2].cells
        cells[0].text = "Overall Offensive Efficiency vs Any Pressure (Gain greater than 3)"
        try:
            cells[1].text = str(round(len(self.d_report.query('Result >= 3 & Pressure_Existence == "Yes"' )) / len(self.d_report.query('Pressure_Existence == "Yes"')) * 100, 2)) + "%" + " | Observations: " + str(len(self.d_report.query('Pressure_Existence == "Yes"')))
        except: 
            cells[1].text = "No Data"

        cells = table.rows[3].cells
        cells[0].text = "Overall Offensive Efficiency vs Edge Pressure (Gain greater than 3)"
        try:
            cells[1].text = str(round(len(self.d_report.query('Result >= 3 & Pressure_Edge == "Yes"' )) / len(self.d_report.query('Pressure_Edge == "Yes"')) * 100, 2)) + "%" + " | Observations: " + str(len(self.d_report.query('Pressure_Edge == "Yes"')))
        except: 
            cells[1].text = "No Data"

        cells = table.rows[4].cells
        cells[0].text = "Overall Offensive Efficiency vs Middle Pressure (Gain greater than 3)"
        try:
            cells[1].text = str(round(len(self.d_report.query('Result >= 3 & Pressure_Center == "Yes"' )) / len(self.d_report.query('Pressure_Center == "Yes"')) * 100, 2)) + "%" + " | Observations: " + str(len(self.d_report.query('Pressure_Center == "Yes"')))
        except:
            cells[1].text = "No Data"

        cells = table.rows[5].cells
        cells[0].text = "Passing"
        table.rows[5].cells[0].paragraphs[0].runs[0].font.bold = True

        cells = table.rows[6].cells
        cells[0].text = "Completion Percentage"
        try:
            cells[1].text = str(round(len(self.d_pass[self.d_pass["Result"] > 0]) / len(self.d_pass) * 100, 2)) + "%" + " | Observations:" + str(len(self.d_pass))
        except: 
            cells[1].text = "No Data"

        cells = table.rows[7].cells
        cells[0].text = "Completion Percentage vs Pressure"
        try:    
            cells[1].text = str(round(len(self.d_pass.query('Result > 0 & Pressure_Existence == "Yes"')) / len(self.d_pass.query('Pressure_Existence == "Yes"')) * 100, 2)) + "%" + " | Observations: " + str(len(self.d_pass.query('Pressure_Existence == "Yes"')))
        except: 
            cells[1].text = "No Data"
            
        cells = table.rows[8].cells
        cells[0].text = "Efficiency vs Zone (Gain greater than 3)"
        try:
            cells[1].text = str(round(len(self.d_pass.query('Result >= 3 & Coverage_Type == "Zone"')) / len(self.d_pass.query('Coverage_Type == "Zone"')) * 100, 2)) + "%" + " | Observations: " + str(len(self.d_pass.query('Coverage_Type == "Zone"')))
        except:
            cells[1].text = "No Data"
            
        cells = table.rows[9].cells
        cells[0].text = "Efficiency vs Man (Gain greater than 3)"
        try:
            cells[1].text = str(round(len(self.d_pass.query('Result >= 3 & Coverage_Type == "Man"')) / len(self.d_pass.query('Coverage_Type == "Man"')) * 100, 2)) + "%" + " | Observations: " + str(len(self.d_pass.query('Coverage_Type == "Man"')))
        except: 
            cells[1].text = "No Data"

        cells = table.rows[10].cells
        cells[0].text = "QB Scramble Rate (Times where QB kept ball on Pass)"
        try:
            cells[1].text = str(round(len(self.d_pass.query('Result > 0 & Pass_Zone == "Non Passing Play"')) / len(self.d_pass) * 100, 2)) + "%" + " | Observations: " + str(len(self.d_pass))
        except: 
            cells[1].text = "No Data"

        cells = table.rows[11].cells
        cells[0].text = "Running"
        table.rows[11].cells[0].paragraphs[0].runs[0].font.bold = True
        cells[1].text = ""

        cells = table.rows[12].cells
        cells[0].text = "Efficiency on Outside Runs (Gain greater than 3)"
        try:
            cells[1].text = str(round(len(self.d_report.query('Run_Type == "Outside" & Result >= 3')) / len(self.d_report.query('Run_Type == "Outside"')) * 100, 2)) + "%" + " | Observations: " + str(len(self.d_report.query('Run_Type == "Outside"')))
        except: 
            cells[1].text = "No Data"
            
        cells = table.rows[13].cells
        cells[0].text = "Efficiency on Inside Runs (Gain greater than 3)"
        try:
            cells[1].text = str(round(len(self.d_report.query('Run_Type == "Inside" & Result >= 3')) / len(self.d_report.query('Run_Type == "Inside"')) * 100, 2)) + "%" + " | Observations: " + str(len(self.d_report.query('Run_Type == "Inside"')))
        except: 
            cells[1].text = "No Data"

        cells = table.rows[14].cells
        cells[0].text = "3rd Down"
        table.rows[14].cells[0].paragraphs[0].runs[0].font.bold = True
        cells[1].text = ""

        cells = table.rows[15].cells
        cells[0].text = "3rd Down Conversion Rate"
        try:
            cells[1].text = str(round(len(self.d_report.query('Down == 3 & Result >= Distance')) / len(self.d_report.query('Down == 3')) * 100, 2)) + "%" + " | Observations: " + str(len(self.d_report.query('Down == 3')))
        except: 
            cells[1].text = "No Data"

        cells = table.rows[16].cells
        cells[0].text = "3rd Down Conversion Rate vs Pressure"
        try:
            cells[1].text = str(round(len(self.d_report.query('Down == 3 & Result >= Distance & Pressure_Existence == "Yes"')) / len(self.d_report.query('Down == 3 & Pressure_Existence == "Yes"')) * 100, 2)) + "%" + " | Observations: " + str(len(self.d_report.query('Down == 3 & Pressure_Existence == "Yes"')))
        except: 
            cells[1].text = "No Data"

        cells = table.rows[17].cells
        cells[0].text = "Sack Rate"
        try:
            cells[1].text = str(round(len(self.d_pass.query('Result < 0 & Pass_Zone == "Non Passing Play"')) / len(self.d_pass) * 100, 2)) + "%" + " | Observations: " + str(len(self.d_pass))
        except: 
            cells[1].text = "No Data"

        self.report.add_page_break()
    
    def Yards(self):
        title = self.report.add_heading('%s Offense - Yardage Breakdown' %self.team_of_interest, 2)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        ## Set Headers
        table = self.report.add_table(8, 5)
        table.style = 'Light Grid'
        headers = table.rows[0].cells
        headers[0].text = "Type"
        headers[1].text = "Total Yards"
        headers[2].text = "Yards Per Game"
        headers[3].text = "Yards Per Drive"
        headers[4].text = "Yards Per Attempt"

        ### Make Bold
        table.rows[0].cells[0].paragraphs[0].runs[0].font.bold = True
        table.rows[0].cells[1].paragraphs[0].runs[0].font.bold = True
        table.rows[0].cells[2].paragraphs[0].runs[0].font.bold = True
        table.rows[0].cells[3].paragraphs[0].runs[0].font.bold = True
        table.rows[0].cells[4].paragraphs[0].runs[0].font.bold = True

        unique_games = len(pd.concat([self.d_report['Team_Name'],self.d_report['Opponent_Name'],self.d_report['Year']]).unique())
        unique_drives = len(self.d_report["Drive"].unique())
        unqiue_plays = len(self.d_report)

        cells = table.rows[1].cells
        cells[0].text = "Total Yards"
        try:
            cells[1].text = str(self.d_report["Result"].sum())
            cells[2].text = '{:0.2f}'.format(self.d_report["Result"].sum() / unique_games)
            cells[3].text = '{:0.2f}'.format(self.d_report["Result"].sum() / unique_drives)
            cells[4].text = '{:0.2f}'.format(self.d_report["Result"].sum() /unqiue_plays)
        except:
            cells[1].text = "No Data"
            cells[2].text = "No Data"
            cells[3].text = "No Data"
            cells[4].text = "No Data"
    
        cells = table.rows[2].cells
        cells[0].text = "Rushing Yards"
        try:   
            cells[1].text = str(self.d_runs["Result"].sum())
            cells[2].text = '{:0.2f}'.format(self.d_runs["Result"].sum() / unique_games)
            cells[3].text = '{:0.2f}'.format(self.d_runs["Result"].sum() / unique_drives)
            cells[4].text = '{:0.2f}'.format(self.d_runs["Result"].sum() / len(self.d_runs)) 
        except: 
            cells[1].text = "No Data"
            cells[2].text = "No Data"
            cells[3].text = "No Data"
            cells[4].text = "No Data"    

        cells = table.rows[3].cells
        cells[0].text = "Inside Run Yards"
        try: 
            cells[1].text = str(self.d_report.query("Play_Type == 'Inside Run'")["Result"].sum())
            cells[2].text = '{:0.2f}'.format(self.d_report.query("Play_Type == 'Inside Run'")["Result"].sum() / unique_games)
            cells[3].text = '{:0.2f}'.format(self.d_report.query("Play_Type == 'Inside Run'")["Result"].sum() / unique_drives)
            cells[4].text = '{:0.2f}'.format(self.d_report.query("Play_Type == 'Inside Run'")["Result"].sum() / len(self.d_report.query("Play_Type == 'Inside Run'"))) 
        except: 
            cells[1].text = "No Data"
            cells[2].text = "No Data"
            cells[3].text = "No Data"
            cells[4].text = "No Data"

        cells = table.rows[4].cells
        cells[0].text = "Outside Run Yards"
        try:    
            cells[1].text = str(self.d_report.query("Play_Type == 'Outside Run'")["Result"].sum())
            cells[2].text = '{:0.2f}'.format(self.d_report.query("Play_Type == 'Outside Run'")["Result"].sum() / unique_games)
            cells[3].text = '{:0.2f}'.format(self.d_report.query("Play_Type == 'Outside Run'")["Result"].sum() / unique_drives)
            cells[4].text = '{:0.2f}'.format(self.d_report.query("Play_Type == 'Outside Run'")["Result"].sum() / len(self.d_report.query("Play_Type == 'Outside Run'"))) 
        except:
            cells[1].text = "No Data"
            cells[2].text = "No Data"
            cells[3].text = "No Data"
            cells[4].text = "No Data"

        cells = table.rows[5].cells
        cells[0].text = "Passing Yards"
        try:    
            cells[1].text = str(self.d_pass["Result"].sum())
            cells[2].text = '{:0.2f}'.format(self.d_pass["Result"].sum() / unique_games)
            cells[3].text = '{:0.2f}'.format(self.d_pass["Result"].sum() / unique_drives)
            cells[4].text = '{:0.2f}'.format(self.d_pass["Result"].sum() / len(self.d_pass))
        except: 
            cells[1].text = "No Data"
            cells[2].text = "No Data"
            cells[3].text = "No Data"
            cells[4].text = "No Data"  

        cells = table.rows[6].cells
        cells[0].text = "Pocket Pass Yards"
        try:    
            cells[1].text = str(self.d_pass.query("Play_Type == 'Pocket Pass'")["Result"].sum())
            cells[2].text = '{:0.2f}'.format(self.d_pass.query("Play_Type == 'Pocket Pass'")["Result"].sum() / unique_games)
            cells[3].text = '{:0.2f}'.format(self.d_pass.query("Play_Type == 'Pocket Pass'")["Result"].sum() / unique_drives)
            cells[4].text = '{:0.2f}'.format(self.d_pass.query("Play_Type == 'Pocket Pass'")["Result"].sum() / len(self.d_pass.query("Play_Type == 'Pocket Pass'")["Result"])) 
        except: 
            cells[1].text = "No Data"
            cells[2].text = "No Data"
            cells[3].text = "No Data"
            cells[4].text = "No Data" 

        cells = table.rows[7].cells
        cells[0].text = "Boot Pass Yards"
        try:    

            cells[1].text = str(self.d_pass.query("Play_Type == 'Boot Pass'")["Result"].sum())
            cells[2].text = '{:0.2f}'.format(self.d_pass.query("Play_Type == 'Boot Pass'")["Result"].sum() / unique_games)
            cells[3].text = '{:0.2f}'.format(self.d_pass.query("Play_Type == 'Boot Pass'")["Result"].sum() / unique_drives)
            cells[4].text = '{:0.2f}'.format(self.d_pass.query("Play_Type == 'Boot Pass'")["Result"].sum() / len(self.d_pass.query("Play_Type == 'Boot Pass'")["Result"])) 
        except: 
            cells[1].text = "No Data"
            cells[2].text = "No Data"
            cells[3].text = "No Data"
            cells[4].text = "No Data" 
        self.report.add_paragraph()

        self.Table(self.d_report.groupby(["Play_Type"]).agg({"Result": ["count", "mean", "median", "std", "max"]}), "Play Type Yardage")
        self.report.add_page_break()

    def PlayType(self):
        ### D Play Type Page 1
        title = self.report.add_heading('%s Offense - Play Type' %self.team_of_interest, 2)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.CrosstabDisplay(pd.crosstab(self.d_report.FieldGroup, self.d_report.Play_Type, margins=True, normalize="index"), "Play Type by Field Position")
        
        ## Swarm Graph
        self.SwarmPlot(self.d_report, "Distance", "Play_Type", "Distance to First Down", "Play Type", 6, 3.5)

        self.d_report["Down"] = self.d_report["Down"].astype(str)
        self.CrosstabDisplay(pd.crosstab(self.d_report.Down, self.d_report.Play_Type, margins=True, normalize="index"), "Play Type by Down")
        self.CrosstabDisplay(pd.crosstab(self.d_report.Hash, self.d_report.Play_Type, margins=True, normalize="index"), "Play Type by Hash")
        self.d_report["Down"] = self.d_report["Down"].astype(int)
        self.report.add_page_break()

        ### D Play Type Page 2
        title = self.report.add_heading('%s Offense - Play Type Continued' %self.team_of_interest, 2)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.CrosstabDisplay(pd.crosstab(self.d_report.Personnel,  self.d_report.Play_Type, margins=True, normalize="index"), "Play Type by Personnel")
        self.CrosstabDisplay(pd.crosstab(self.d_report.Formation, self.d_report.Play_Type, margins=True, normalize="index"), "Play Type by Formation")
        self.report.add_page_break()

    def BallCarrier(self):
        ### Ball Carrier
        self.d_runs["Result_BallCarrier"] = self.d_runs["Result_BallCarrier"].apply(lambda x: str(x))
        self.d_runs["Down"] = self.d_runs["Down"].apply(lambda x: str(x))

        title = self.report.add_heading('%s Offense - Ball Carrier' %self.team_of_interest, 2)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.Table(self.d_runs.groupby("Result_BallCarrier").agg({"Result": ["count", "sum", "mean", "max", "median"]}), "Ball Carrier Stats")
        self.CrosstabDisplay(pd.crosstab(self.d_runs.Personnel,  self.d_runs.Result_BallCarrier, margins=True, normalize="index"), "Ball Carrier by Personnel")
        self.CrosstabDisplay(pd.crosstab(self.d_runs.Formation, self.d_runs.Result_BallCarrier, margins=True, normalize="index"), "Ball Carrier by Formation")

        # TO-DO: Add Ball Carrier Statistics
        self.CrosstabDisplay(pd.crosstab(self.d_runs.Play_Type,  self.d_runs.Result_BallCarrier, margins=True, normalize="index"), "Ball Carrier by Play_Type")
        self.report.add_page_break()

        ### Ball Carrier continued
        title = self.report.add_heading('%s Offense - Ball Carrier' %self.team_of_interest, 2)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.CrosstabDisplay(pd.crosstab(self.d_runs.Down,  self.d_runs.Result_BallCarrier, margins=True, normalize="index"), "Ball Carrier by Down")
        self.SwarmPlot(self.d_runs, "Distance", "Result_BallCarrier", "Distance to First Down", "Ball Carrier", 6, 3.5)
        self.SwarmPlot(self.d_runs, "Yard", "Result_BallCarrier", "Field Position", "Ball Carrier", 6, 3.5)
        self.report.add_page_break()

    def Down(self):
        ### D 1st Down Analysis 
        title = self.report.add_heading('%s Offense - Situational Analysis 1st Down' %self.team_of_interest, 2)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.report.add_heading("1st Down Play Type by Distance & Hash", 3)
        self.Mapper(pd.crosstab(self.d_1_down.Distance_Hash, self.d_1_down.Play_Type, normalize="index"), "Play Type", "Hash Mark and Distance", 3.6, 2.5)
        self.report.add_heading("1st Down Play Type by Field Position", 3)
        self.Mapper(pd.crosstab(self.d_1_down.FieldGroup, self.d_1_down.Play_Type, normalize="index"), "Play Type", "Field Position", 3.6, 2.91)
        self.CrosstabDisplay(pd.crosstab(self.d_1_down.Personnel, self.d_1_down.Play_Type, margins=True, normalize="index"), "Play Type by Personnel")
        self.report.add_page_break()

        ### D 2nd Down Analysis 
        self.report.add_heading('%s Offense - Situational Analysis 2nd Down' %self.team_of_interest, 2)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.report.add_heading("2nd Down Play Type by Distance & Hash", 3)
        self.Mapper(pd.crosstab(self.d_2_down.Distance_Hash, self.d_2_down.Play_Type, normalize="index"), "Play Type", "Hash Mark and Distance", 3.6, 2.5)
        self.report.add_heading("2nd Down Play Type by Field Position", 3)
        self.Mapper(pd.crosstab(self.d_2_down.FieldGroup, self.d_2_down.Play_Type, normalize="index"), "Play Type", "Field Position", 3.6, 2.91)
        self.CrosstabDisplay(pd.crosstab(self.d_2_down.Personnel, self.d_2_down.Play_Type, margins=True, normalize="index"), "Play Type by Personnel")
        self.report.add_page_break()

        ### D 3rd Down Analysis 
        self.report.add_heading('%s Offense - Situational Analysis 3rd Down' %self.team_of_interest, 2)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.report.add_heading("3rd Down Play Type by Distance & Hash", 3)
        self.Mapper(pd.crosstab(self.d_3_down.Distance_Hash, self.d_3_down.Play_Type, normalize="index"), "Play Type", "Hash Mark and Distance", 3.6, 2.5)
        self.report.add_heading("3rd Down Play Type by Field Position", 3)
        self.Mapper(pd.crosstab(self.d_3_down.FieldGroup, self.d_3_down.Play_Type, normalize="index"), "Play Type", "Field Position", 3.6, 2.91)
        self.CrosstabDisplay(pd.crosstab(self.d_3_down.Personnel, self.d_3_down.Play_Type, margins=True, normalize="index"), "Play Type by Personnel")
        self.report.add_page_break()

    def IdealCoverage(self):
        title = self.report.add_heading('%s Offense - Ideal Passing Coverage' %self.team_of_interest, 2)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        self.Table(self.d_pass.groupby(["Down", "Coverage"]).agg({"Result": ["count", "mean", "median", "std"]}), "Coverage Detailed by Down")
        self.Table(self.d_pass.groupby(["DownGroup", "Coverage"]).agg({"Result": ["count", "mean", "median", "std"]}), "Coverage Detailed by Down Distance")
        self.Table(self.d_pass.groupby(["FieldGroup", "Coverage"]).agg({"Result": ["count", "mean", "median", "std"]}), "Coverage Detailed by Down Field Position")
        self.report.add_page_break()

    def PassZone(self):
        ### D Pass Zones
        title = self.report.add_heading('%s Offense - Pass Zones' %self.team_of_interest, 2)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.d_passing_zones["Down"] = self.d_passing_zones["Down"].astype(str)
        self.CrosstabDisplay(pd.crosstab(self.d_passing_zones.Down, self.d_passing_zones.Pass_Zone, margins=True, normalize="index"), "Passing Zone by Down")
        self.CrosstabDisplay(pd.crosstab(self.d_passing_zones.DownGroup, self.d_passing_zones.Pass_Zone, margins=True, normalize="index"), "Passing Zone by Distance")
        self.SwarmPlot(self.d_passing_zones, "Yard", "Pass_Zone", "Field Position", "Passing Zone", 6, 3.5)
        self.CrosstabDisplay(pd.crosstab(self.d_passing_zones.Personnel, self.d_passing_zones.Pass_Zone, margins=True, normalize="index"), "Pass Zone by Personnel")
        self.report.add_page_break()

    def Targets(self):
        title = self.report.add_heading('%s Offense - Pass Zone Targets' %self.team_of_interest, 2)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.d_passing_zones["Result_BallCarrier"] = self.d_passing_zones["Result_BallCarrier"].astype(str)


        self.CrosstabDisplay(pd.crosstab(self.d_passing_zones.Pass_Zone, self.d_passing_zones.Coverage, margins=True, normalize="index"), "Coverage by Passzone")
        self.CrosstabDisplay(pd.crosstab(self.d_passing_zones.Pass_Zone, self.d_passing_zones.Result_BallCarrier, margins=True, normalize="index"), "Targets by Passzone")
        self.Table(self.d_passing_zones.groupby("Result_BallCarrier").agg({"Result": ["count", "sum", "mean", "median", "max"]}), "Reciever Yardage")
        self.report.add_page_break()

    def Redzone(self):
        ### D Redzone
        title = self.report.add_heading('%s Offense - Redzone' %self.team_of_interest, 2)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title = self.report.add_heading('Redzone - Play Type', 3)
        self.CrosstabDisplay(pd.crosstab(self.d_redzone.Redzone, self.d_redzone.Play_Type, margins=True, normalize="index"), "Play Type by Redzone Position")
        self.d_redzone["Down"] = self.d_redzone["Down"].astype(str)
        self.CrosstabDisplay(pd.crosstab(self.d_redzone.Down, self.d_redzone.Play_Type, margins=True, normalize="index"), "Play Type by Down")
        self.CrosstabDisplay(pd.crosstab(self.d_redzone.Personnel, self.d_redzone.Play_Type, margins=True, normalize="index"), "Play Type by Personnel")
        self.CrosstabDisplay(pd.crosstab(self.d_redzone.Formation, self.d_redzone.Play_Type, margins=True, normalize="index"), "Play Type by Formation")
        self.report.add_page_break()

    def Boundary(self):
        title = self.report.add_heading('%s Offense - Boundary' %self.team_of_interest, 2)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title = self.report.add_heading('Analyzing Left and Right Hash Plays', 4)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Play Type by Down
        self.d_boundary["Down"] = self.d_boundary["Down"].astype(str)
        self.CrosstabDisplay(pd.crosstab(self.d_boundary.Down, self.d_boundary.Play_Type, margins=True, normalize="index"), "Play Type by Down")
        self.CrosstabDisplay(pd.crosstab(self.d_boundary.Down, self.d_boundary.Boundary, margins=True, normalize="index"), "Boundary by Down")
        # Play Type into Boundary
        self.CrosstabDisplay(pd.crosstab(self.d_boundary.Boundary, self.d_boundary.Play_Type, margins=True, normalize="index"), "Boundary by Play Type")
        # Plays to Strength
        self.CrosstabDisplay(pd.crosstab(self.d_boundary.Boundary, self.d_boundary.Hash, margins=True, normalize="index"), "Boundary by Hash")
        self.report.add_page_break()
        
    def Temporal(self):
        pass

def run_pre_report(input: list, team_of_interest: list, team_code: str, type: str):
    warnings.filterwarnings("ignore")
    return PregameReport(input, team_of_interest, team_code, type).output_path
