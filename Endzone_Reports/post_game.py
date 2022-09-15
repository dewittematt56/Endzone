import sys
from graph_utils import Graph_Utils
from docx import Document
import time
import logging
import os
root_folder = os.path.abspath(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
sys.path.append(root_folder)
from datetime import datetime
from sqlalchemy import create_engine
from sqlalchemy.sql import functions 
from sqlalchemy.orm import sessionmaker
from Endzone_Database.classes import Game
import pandas as pd
from docx.enum.text import WD_ALIGN_PARAGRAPH

class PostgameReport():
    def __init__(self, input, teamCode):
        self.jobId = str(int(time.time()))
        self.input = input
        self.team_code = teamCode
        self.report = Document()

        logging.basicConfig(filename = os.path.join(os.getcwd(), "Endzone_Reports/Logs/PreGame_Log.txt"), level=logging.INFO)
        logging.info("Job | " + self.jobId +  " Start | " + str(self.input) + "\t" + str(datetime.fromtimestamp(time.time()).isoformat()))
        
        self.data = self.get_data()
        self.data = self.prep_data()
        self.report_manager()
    
    def get_data(self):
        db_engine = create_engine("postgresql://endzone:Hercules22!@172.104.25.168/Endzone")
        Session = sessionmaker(db_engine)
        session = Session()
        data = pd.read_sql(session.query(Game).filter(Game.Owner_Team_Code == self.team_code).filter(functions.concat(Game.Team_Name, "_", Game.Opponent_Name, "_", Game.Year).in_(self.input)).statement, db_engine)
        if len(data) > 0:
            return data
        else:
            logging.info("\t Job " + self.jobId +  " Error | \t" + str(self.input) + "\t" + str(datetime.fromtimestamp(time.time()).isoformat()))
            logging.error("No Data found")
            return

    def prep_data(self):
        self.data["DownGroup"] = ""
        self.data["FieldGroup"] = ""
        self.data["Pressure_Stunt"] = ""
        self.data["Pressure_Center"] = ""
        self.data["Pressure_Edge"] = ""
        self.data["Pressure_Existence"] = ""
        self.data["Number_Rushers"] = 0
        self.data["Coverage_Type"] = ""
        self.data["Run_Type"] = ""
        self.data["Personal"] = ""
        for index, row in self.data.iterrows():
            try:
                self.data.at[index, "Number_Rushers"] = str(self.data.at[index, "D_Formation"])[0:1]
                if self.data.at[index, "Distance"] <= 3:
                    self.data.at[index, "DownGroup"] = "Short"
                elif self.data.at[index, "Distance"] > 3 and self.data.at[index, "Distance"] <= 6:
                    self.data.at[index, "DownGroup"] = "Medium"
                elif self.data.at[index, "Distance"] > 6:
                    self.data.at[index, "DownGroup"] = "Long"
                if self.data.at[index, "Yard"] <= 33:
                    self.data.at[index, "FieldGroup"] = "Backed Up"
                elif self.data.at[index, "Yard"] > 33 and self.data.at[index, "Yard"] <= 66:
                    self.data.at[index, "FieldGroup"] = "Midfield"
                elif self.data.at[index, "Yard"] > 66:
                    self.data.at[index, "FieldGroup"] = "Scoring Position"
                if int(self.data.at[index, "Pressure_Middle"]) > 0:
                    self.data.at[index, "Pressure_Center"] = "Yes"
                    self.data.at[index, "Number_Rushers"] = int(self.data.at[index, "Number_Rushers"]) + int(self.data.at[index, "Pressure_Middle"])
                else:
                    self.data.at[index, "Pressure_Center"] = "No"
                if int(self.data.at[index, "Pressure_Right"]) > 0:
                    self.data.at[index, "Number_Rushers"] = int(self.data.at[index, "Number_Rushers"]) + int(self.data.at[index, "Pressure_Right"])
                if int(self.data.at[index, "Pressure_Left"]) > 0:
                    self.data.at[index, "Number_Rushers"] = int(self.data.at[index, "Number_Rushers"]) + int(self.data.at[index, "Pressure_Left"])
                if int(self.data.at[index, "Pressure_Left"]) > 0 or int(self.data.at[index, "Pressure_Right"]) > 0:
                    self.data.at[index, "Pressure_Edge"] = "Yes"
                else:
                    self.data.at[index, "Pressure_Edge"] = "No"
                if self.data.at[index, "Number_Rushers"] != str(self.data.at[index, "D_Formation"])[0:1]:
                    self.data.at[index, "Pressure_Existence"] = "Yes"
                else:
                    self.data.at[index, "Pressure_Existence"] = "No"       
                if "Zone" in str(self.data.at[index, "Coverage"]):
                    self.data.at[index, "Coverage_Type"] = "Zone"
                elif "Man" in str(self.data.at[index, "Coverage"]):
                    self.data.at[index, "Coverage_Type"] = "Man"
                else:
                    self.data.at[index, "Coverage_Type"] = "Unknown" 
                if "Outside" in self.data.at[index, "Play_Type"]:
                    self.data.at[index, "Run_Type"] = "Outside"
                elif "Inside" in self.data.at[index, "Play_Type"]:
                    self.data.at[index, "Run_Type"] = "Inside"
                else:
                    self.data.at[index, "Run_Type"] = "Null" 
            except Exception as e:
                logging.error(e)
                continue
        self.data["FieldSpot"] = self.data["Hash"] + "-" + self.data["FieldGroup"]
        return self.data

    def report_manager(self):

            ### Title Page
            title = self.report.add_heading("EVFB Analytics Postgame Analysis")
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for team in self.data.Possession.unique():
                data = self.data[(self.data["Possession"] == team)]
                Content(data, self.report, team)
            
            self.output_path = os.path.dirname(__file__) + "/Outputs/EVFB_PostGame_Report_%s.docx" %self.jobId
            self.report.save(self.output_path)


class Content(Graph_Utils):
    def __init__(self, data, report, team):
        self.data = data
        self.report = report
        self.team = team

        self.report.add_heading(team + " Offense", 2)
        self.create_content()

    def create_content(self):
        self.total_yards()
        self.passing()
        self.targets()
        self.rushing()
        self.report.add_page_break()

    def total_yards(self):
        self.report.add_heading("Total Yards")

        table = self.report.add_table(15, 4)
        table.style = 'Light Grid'
        headers = table.rows[0].cells
        headers[0].text = "Type"
        headers[1].text = "Total"
        headers[2].text = "Per Drive"
        headers[3].text = "Per Attempt"

        ### Make Bold
        table.rows[0].cells[0].paragraphs[0].runs[0].font.bold = True
        table.rows[0].cells[1].paragraphs[0].runs[0].font.bold = True
        table.rows[0].cells[2].paragraphs[0].runs[0].font.bold = True
        table.rows[0].cells[3].paragraphs[0].runs[0].font.bold = True

        unique_drives = len(self.data["Drive"].unique())
        unqiue_plays = len(self.data)

        cells = table.rows[1].cells
        cells[0].text = "Total Yards"
        try:
            cells[1].text = str(self.data["Result"].sum())
            cells[2].text = '{:0.2f}'.format(self.data["Result"].sum() / unique_drives)
            cells[3].text = '{:0.2f}'.format(self.data["Result"].sum() /unqiue_plays)
        except:
            cells[1].text = "No Data"
            cells[2].text = "No Data"
            cells[3].text = "No Data"
    
        cells = table.rows[2].cells
        cells[0].text = "Total Rushing Yards"
        try:   
            cells[1].text = str(self.data.query("Play_Type == 'Inside Run' or Play_Type == 'Outside Run'")["Result"].sum() - self.data.query("(Play_Type == 'Pocket Pass' or Play_Type == 'Boot Pass') and Result < 0")["Result"].sum()) 
            cells[2].text = '{:0.2f}'.format(self.data.query("Play_Type == 'Inside Run' or Play_Type == 'Outside Run'")["Result"].sum() / unique_drives)
            cells[3].text = '{:0.2f}'.format(self.data.query("Play_Type == 'Inside Run' or Play_Type == 'Outside Run'")["Result"].sum() /len(self.data.query("Play_Type == 'Inside Run' or Play_Type == 'Outside Run'"))) 
        except: 
            cells[1].text = "No Data"
            cells[2].text = "No Data"
            cells[3].text = "No Data"

        cells = table.rows[3].cells
        cells[0].text = "Inside Run Yards"
        try: 
            cells[1].text = str(self.data.query("Play_Type == 'Inside Run'")["Result"].sum())
            cells[2].text = '{:0.2f}'.format(self.data.query("Play_Type == 'Inside Run'")["Result"].sum() / unique_drives)
            cells[3].text = '{:0.2f}'.format(self.data.query("Play_Type == 'Inside Run'")["Result"].sum() / len(self.data.query("Play_Type == 'Inside Run'"))) 
        except: 
            cells[1].text = "No Data"
            cells[2].text = "No Data"
            cells[3].text = "No Data"

        cells = table.rows[4].cells
        cells[0].text = "Outside Run Yards"
        try:    
            cells[1].text = str(self.data.query("Play_Type == 'Outside Run'")["Result"].sum())
            cells[2].text = '{:0.2f}'.format(self.data.query("Play_Type == 'Outside Run'")["Result"].sum() / unique_drives)
            cells[3].text = '{:0.2f}'.format(self.data.query("Play_Type == 'Outside Run'")["Result"].sum() / len(self.data.query("Play_Type == 'Outside Run'"))) 
        except:
            cells[1].text = "No Data"
            cells[2].text = "No Data"
            cells[3].text = "No Data"

        cells = table.rows[5].cells
        cells[0].text = "Total Passing Yards"
        try:    
            cells[1].text = str(self.data.query("Play_Type == 'Pocket Pass' or Play_Type == 'Boot Pass'")["Result"].sum())
            cells[2].text = '{:0.2f}'.format(self.data.query("Play_Type == 'Pocket Pass' or Play_Type == 'Boot Pass'")["Result"].sum() / unique_drives)
            cells[3].text = '{:0.2f}'.format(self.data.query("Play_Type == 'Pocket Pass' or Play_Type == 'Boot Pass'")["Result"].sum() / len(self.data.query("Play_Type == 'Pocket Pass' | Play_Type == 'Boot Pass'")))
        except: 
            cells[1].text = "No Data"
            cells[2].text = "No Data"
            cells[3].text = "No Data"

        cells = table.rows[6].cells
        cells[0].text = "Pocket Pass Yards"
        try:    
            cells[1].text = str(self.data.query("Play_Type == 'Pocket Pass'")["Result"].sum())
            cells[2].text = '{:0.2f}'.format(self.data.query("Play_Type == 'Pocket Pass'")["Result"].sum() / unique_drives)
            cells[3].text = '{:0.2f}'.format(self.data.query("Play_Type == 'Pocket Pass'")["Result"].sum() / len(self.data.query("Play_Type == 'Pocket Pass'"))) 
        except: 
            cells[1].text = "No Data"
            cells[2].text = "No Data"
            cells[3].text = "No Data"

        cells = table.rows[7].cells
        cells[0].text = "Boot Pass Yards"
        try:    

            cells[1].text = str(self.data.query("Play_Type == 'Boot Pass'")["Result"].sum())
            cells[2].text = '{:0.2f}'.format(self.data.query("Play_Type == 'Boot Pass'")["Result"].sum() / unique_drives)
            cells[3].text = '{:0.2f}'.format(self.data.query("Play_Type == 'Boot Pass'")["Result"].sum() / len(self.data.query("Play_Type == 'Boot Pass'"))) 
        except: 
            cells[1].text = "No Data"
            cells[2].text = "No Data"
            cells[3].text = "No Data"

        cells = table.rows[8].cells
        cells[0].text = "Number of Plays"
        table.rows[8].cells[0].paragraphs[0].runs[0].font.bold = True

        cells = table.rows[9].cells
        cells[0].text = "Play's Ran"
        try:    

            cells[1].text = str(len(self.data))
            cells[2].text = '{:0.2f}'.format(len(self.data) / unique_drives)
            cells[3].text = ' '
        except: 
            cells[1].text = "No Data"
            cells[2].text = "No Data"
            cells[3].text = ' '    

        cells = table.rows[10].cells
        cells[0].text = "Scoring -- ESTIMATE"
        table.rows[10].cells[0].paragraphs[0].runs[0].font.bold = True

        cells = table.rows[11].cells
        cells[0].text = "Inside Run Touchdowns"
        try:    

            cells[1].text = str(len(self.data.query))
            cells[2].text = '{:0.2f}'.format(len(self.data.query) / unique_drives)
            cells[3].text = ' '
        except: 
            cells[1].text = str(len(self.data.query("Play_Type == 'Inside Run' and Result + Yard > 100")))
            cells[2].text = ' '
            cells[3].text = ' ' 

        cells = table.rows[12].cells
        cells[0].text = "Outside Run Touchdowns"
        try:    

            cells[1].text = str(len(self.data.query("Play_Type == 'Outside Run' and Result + Yard > 100")))
            cells[2].text = ''
            cells[3].text = ' '
        except: 
            cells[1].text = "No Data"
            cells[2].text = ' '
            cells[3].text = ' ' 

        cells = table.rows[13].cells
        cells[0].text = "Pocket Pass Touchdowns"
        try:    

            cells[1].text = str(len(self.data.query("Play_Type == 'Pocket Pass' and Result + Yard > 100")))
            cells[2].text = ' '
            cells[3].text = ' '
        except: 
            cells[1].text = "No Data"
            cells[2].text = ' '
            cells[3].text = ' ' 

        cells = table.rows[14].cells
        cells[0].text = "Boot Pass Touchdown's"
        try:    

            cells[1].text = str(len(self.data.query("Play_Type == 'Boot Pass' and Result + Yard > 100")))
            cells[2].text = ' '
            cells[3].text = ' '
        except: 
            cells[1].text = "No Data"
            cells[2].text = ' '
            cells[3].text = ' ' 

        self.report.add_paragraph()

    

    def passing(self):
        table = self.report.add_table(2, 10)

        table.style = 'Light Grid'
        headers = table.rows[0].cells
        headers[0].text = "Attempts"
        headers[1].text = "Completions"
        headers[2].text = "Yards"
        headers[3].text = "Completion %"
        headers[4].text = "TD's"
        headers[5].text = "INT's --IN DEV"
        headers[6].text = "College Rating -- IN DEV"
        headers[7].text = "NFL Rating -- IN DEV"
        headers[8].text = "Yards/Attempt"
        headers[9].text = "Yards/Completion"

        try:
            cells = table.rows[1].cells
            cells[0].text = str(len(self.data.query("(Play_Type == 'Boot Pass' or Play_Type == 'Pocket Pass') and Pass_Zone != 'Non Passing Play'")))
            cells[1].text = str(len(self.data.query("(Play_Type == 'Boot Pass' or Play_Type == 'Pocket Pass') and Pass_Zone != 'Non Passing Play' and Result != 0")))
            cells[2].text = str(self.data.query("(Play_Type == 'Boot Pass' or Play_Type == 'Pocket Pass')")["Result"].sum())
            cells[3].text = str(round(int(cells[1].text) / int(cells[0].text), 2) * 100) + "%"
            cells[4].text = ' '
            cells[5].text = ' '
            cells[6].text = ' '
            cells[7].text = ' '
            cells[8].text = str(round(int(cells[2].text) / int(cells[0].text), 2))
            cells[9].text = str(round(int(cells[2].text) / int(cells[1].text), 2))

        except:
            cells[0].text = ' '
        self.report.add_paragraph()

    def targets(self):
        self.Table(self.data.query("Play_Type == 'Boot Pass' or Play_Type == 'Pocket Pass'").groupby(["Result_BallCarrier"]).agg({"Result": ["count", "sum", "mean", "median", "max"]}), "Reciever Stats")
        self.Table(self.data.query("Play_Type == 'Boot Pass' or Play_Type == 'Pocket Pass'").groupby(["Pass_Zone"]).agg({"Result": ["count", "sum", "mean", "median", "max"]}), "Pass Zone Stats")
    def rushing(self):
        self.Table(self.data.query("Play_Type == 'Inside Run' or Play_Type == 'Outside Run'").groupby(["Result_BallCarrier"]).agg({"Result": ["count", "sum", "mean", "median", "max"]}), "Ball Carrier Stats")

def run_report(input: list, team_code: str):
    return PostgameReport(input, team_code).output_path
run_report(["Eastview_LN_2022"], "EVHS22")