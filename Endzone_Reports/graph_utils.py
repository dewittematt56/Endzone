from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
from io import BytesIO
import matplotlib.pyplot as plt
import seaborn as sns
from warnings import filterwarnings
import matplotlib

class Graph_Utils():
    def __init__(self):
        filterwarnings("ignore")
        plt.switch_backend('agg')
        matplotlib.font_manager._load_fontmanager(try_read_cache=False)

    def AddGraphToPage(self, w_size, h_size):
        table = self.report.add_table(1, 1)
        p = table.rows[0].cells[0].add_paragraph()
        p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run()
        r.add_picture(self.memory_space, width=Inches(w_size), height=Inches(h_size))
        plt.clf()
        self.memory_space = BytesIO()
        table.style = 'Light Grid'
        self.report.add_paragraph()

    def BarGraph(self, data, xCol: str, byCol: str, Xaxis: str, Yaxis: str, w_size, h_size):
        sns.countplot(data = data, x = xCol, hue = byCol)
        plt.xlabel(Xaxis)
        plt.ylabel(Yaxis)
        mng = plt.get_current_fig_manager()
        mng.full_screen_toggle()
        plt.savefig(self.memory_space, bbox_inches='tight')
        self.AddGraphToPage(w_size, h_size)
        return

    def LineGraph(data, xCol: str, yCol:str, ):
        sns.lineplot(data = data, x = xCol, y = yCol)
        pass

    def SwarmPlot(self, data, xCol: str, yCol: str, Xaxis: str, Yaxis: str, w_size, h_size):
        sns.swarmplot(data = data, x = xCol, y = yCol)
        plt.xlabel(Xaxis)
        plt.ylabel(Yaxis)
        mng = plt.get_current_fig_manager()
        mng.full_screen_toggle()
        plt.savefig(self.memory_space, bbox_inches='tight')
        self.AddGraphToPage(w_size, h_size)
        return

    def Mapper(self, dataframe, Xaxis, Yaxis, w_size, h_size):
        sns.heatmap(dataframe, cmap = "RdYlBu_r", center = .5, linewidths = 0.5, annot = True)
        plt.xlabel(Xaxis)
        plt.ylabel(Yaxis)
        mng = plt.get_current_fig_manager()
        mng.full_screen_toggle()
        plt.savefig(self.memory_space, bbox_inches='tight')
        self.AddGraphToPage(w_size, h_size)
        return

    def CrosstabDisplay(self, df, title):
        table = self.report.add_table(df.shape[0]+1, df.shape[1] + 1)
        table.cell(0, 0).text = title
        table.rows[0].cells[0].paragraphs[0].runs[0].font.bold = True
        j = 0
        while j < len(df.columns):
            table.cell(0, j + 1).text = df.columns[j]
            j = j + 1
        j = 0
        while j < len(df.index):
            table.cell(j + 1, 0).text = df.index[j]
            j = j + 1
        for i in range(df.shape[0]):
            for j in range(df.shape[-1]):
                table.cell(i + 1, j + 1).text = str(round(df.values[i,j] * 100, 1)) + "%"
        table.style = 'Light Grid'
        self.report.add_paragraph()

    def Table(self, df, title):
        table = self.report.add_table(df.shape[0]+1, df.shape[1] + 1)
        table.cell(0, 0).text = title
        table.rows[0].cells[0].paragraphs[0].runs[0].font.bold = True
        j = 0

        while j < len(df.columns):
            table.cell(0, j + 1).text = df.columns[j]
            j = j + 1

        j = 0
        while j < len(df):
            table.cell(j + 1, 0).text = str(df.index[j])
            j += 1

        for i in range(df.shape[0]):
            for j in range(df.shape[-1]):
                table.cell(i + 1, j + 1).text = '{:0.2f}'.format(df.values[i, j])
        table.style = 'Light Grid'
        self.report.add_paragraph()
        