import sys
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT

import time
import logging
import os
import warnings
from datetime import datetime
from sqlalchemy import create_engine
from sqlalchemy.sql import functions 
from sqlalchemy.orm import sessionmaker
from Endzone_Database.classes import Game
import pandas as pd


class Thunderbolt():
    def __init__(self, teamCode: str, jobType: str, input: list):
        self.jobId = str(int(time.time())) + "_thunderbolt" 
        self.jobType = jobType
        self.team_code = teamCode
        self.document = Document()
        self.input = input
        self.data = self.get_data()
        self.data = self.prep_data()
        self.PageManager()
    def get_data(self):
        db_engine = create_engine("postgresql://endzone:Hercules22!@172.104.25.168/Endzone")
        Session = sessionmaker(db_engine)
        session = Session()
        data = pd.read_sql(session.query(Game).filter(Game.Owner_Team_Code == self.team_code).filter(functions.concat(Game.Team_Name, "_", Game.Opponent_Name, "_", Game.Year).in_(self.input)).statement, db_engine)
        if len(data) > 0:
            return data
        else:
            logging.info("\t Job " + self.jobId +  " Error | \t" + str(self.input) + "\t" + str(datetime.fromtimestamp(time.time()).isoformat()))
            logging.error("No Data found")
            return

    def prep_data(self):
        self.data["DownGroup"] = ""
        self.data["FieldGroup"] = ""
        self.data["Pressure_Stunt"] = ""
        self.data["Pressure_Center"] = ""
        self.data["Pressure_Edge"] = ""
        self.data["Pressure_Existence"] = ""
        self.data["Number_Rushers"] = 0
        self.data["Coverage_Type"] = ""
        self.data["Run_Type"] = ""
        self.data["Personal"] = ""
        for index, row in self.data.iterrows():
            try:
                self.data.at[index, "Number_Rushers"] = str(self.data.at[index, "D_Formation"])[0:1]
                if self.data.at[index, "Distance"] <= 3:
                    self.data.at[index, "DownGroup"] = "Short"
                elif self.data.at[index, "Distance"] > 3 and self.data.at[index, "Distance"] <= 6:
                    self.data.at[index, "DownGroup"] = "Medium"
                elif self.data.at[index, "Distance"] > 6:
                    self.data.at[index, "DownGroup"] = "Long"
                if self.data.at[index, "Yard"] <= 33:
                    self.data.at[index, "FieldGroup"] = "Backed Up"
                elif self.data.at[index, "Yard"] > 33 and self.data.at[index, "Yard"] <= 66:
                    self.data.at[index, "FieldGroup"] = "Midfield"
                elif self.data.at[index, "Yard"] > 66:
                    self.data.at[index, "FieldGroup"] = "Scoring Position"
                if int(self.data.at[index, "Pressure_Middle"]) > 0:
                    self.data.at[index, "Pressure_Center"] = "Yes"
                    self.data.at[index, "Number_Rushers"] = int(self.data.at[index, "Number_Rushers"]) + int(self.data.at[index, "Pressure_Middle"])
                else:
                    self.data.at[index, "Pressure_Center"] = "No"
                if int(self.data.at[index, "Pressure_Right"]) > 0:
                    self.data.at[index, "Number_Rushers"] = int(self.data.at[index, "Number_Rushers"]) + int(self.data.at[index, "Pressure_Right"])
                if int(self.data.at[index, "Pressure_Left"]) > 0:
                    self.data.at[index, "Number_Rushers"] = int(self.data.at[index, "Number_Rushers"]) + int(self.data.at[index, "Pressure_Left"])
                if int(self.data.at[index, "Pressure_Left"]) > 0 or int(self.data.at[index, "Pressure_Right"]) > 0:
                    self.data.at[index, "Pressure_Edge"] = "Yes"
                else:
                    self.data.at[index, "Pressure_Edge"] = "No"
                if self.data.at[index, "Number_Rushers"] != str(self.data.at[index, "D_Formation"])[0:1]:
                    self.data.at[index, "Pressure_Existence"] = "Yes"
                else:
                    self.data.at[index, "Pressure_Existence"] = "No"       
                if "Zone" in str(self.data.at[index, "Coverage"]):
                    self.data.at[index, "Coverage_Type"] = "Zone"
                elif "Man" in str(self.data.at[index, "Coverage"]):
                    self.data.at[index, "Coverage_Type"] = "Man"
                else:
                    self.data.at[index, "Coverage_Type"] = "Unknown" 
                if "Outside" in self.data.at[index, "Play_Type"]:
                    self.data.at[index, "Run_Type"] = "Outside"
                elif "Inside" in self.data.at[index, "Play_Type"]:
                    self.data.at[index, "Run_Type"] = "Inside"
                else:
                    self.data.at[index, "Run_Type"] = "Null" 
                # Flip Perspective to Defense
                if self.jobType == "Defense":
                    if self.data.at[index, "Hash"] == "Left": self.data.at[index, "Hash"] = "Right"
                    elif self.data.at[index, "Hash"] == "Right": self.data.at[index, "Hash"] = "Left"

            except Exception as e:
                logging.error(e)
                continue
        self.data["FieldSpot"] = self.data["Hash"] + "-" + self.data["FieldGroup"]
        return self.data

    def PageManager(self):
        self.section = self.document.add_section(WD_ORIENT.PORTRAIT)
        self.section.left_margin = Inches(0.25)
        self.section.right_margin = Inches(0.25)
        self.section.top_margin = Inches(0.25)
        self.section.bottom_margin = Inches(0.25)

        title = self.document.add_heading("Endzone - Thunderbolt I - %s" %self.jobType)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        self.DefenseSheet()


        self.output_path = os.path.dirname(__file__) + "/Endzone_Thunderbolt.docx"
        print(self.output_path)
        self.document.save(self.output_path)

    def DefenseSheet(self):
        table = self.document.add_table(4, 4, 'Light Grid')

        headers = table.rows[0].cells
        headers[0].text = "Backed Up"
        headers[1].text = "Left"
        headers[2].text = "Middle"
        headers[3].text = "Right"
if __name__ == '__main__':
    job = Thunderbolt("EVHS22", "Defense", ["Eastview_Edina_2022", "Eastview_East-Ridge_2022"])